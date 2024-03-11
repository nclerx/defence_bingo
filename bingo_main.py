import random
import os
import numpy as np
from docx import Document
from docx.shared import Inches, Pt, Cm
from docx.enum.text import *
from docx.enum.table import *

words_all = ['elephant', 'turtleneck', 'skeleton', 'vacuum cleaner', 'humanoid', 'round', 'grenadine', 'redundant',
             'coffee', 'stereotype', 'curling', 'spoiler', 'jazz', 'sock', 'wave', 'volcano', 'Birkenstock',
             'helicopter',
             'peanut butter', 'swimming pool', 'butterfly', 'curtain', 'wrestle', 'alphabet', 'magic', 'shampoo',
             'donkey',
             'jewel', 'anthropocentric', 'tale', 'sofa', 'popcorn', 'digress', 'helmet', 'society']
words = ['elephant', 'turtleneck', 'vacuum cleaner', 'humanoid', 'round', 'grenadine', 'redundant',
         'coffee', 'stereotype', 'curling', 'jazz', 'sock', 'volcano', 'Birkenstock', 'helicopter', 'wrestle',
         'alphabet', 'magic', 'shampoo', 'donkey', 'anthropocentric', 'tale', 'sofa', 'popcorn']


# def get_cardnum():
#     cardnum = int(input('Enter the number of Bingo cards to generate: '))
#     print('%s Bingo-cards will be created in 1 word-file' % cardnum)
#
#     return cardnum


def generate_wordorder(inp):
    random.shuffle(inp)
    row1 = inp[:5]
    row2 = inp[5:10]
    row3 = inp[10:14]
    row3.insert(2, '*')
    row4 = inp[14:19]
    row5 = inp[19:]

    return [row1, row2, row3, row4, row5]


def create_bingogrid():
    table = document.add_table(rows=5, cols=5, style="Table Grid")
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    wrds = generate_wordorder(words)
    for row in np.arange(5):
        table.rows[row].height_rule = WD_ROW_HEIGHT_RULE.EXACTLY
        table.rows[row].height = Cm(2)
        p = table.rows[row]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        rw = table.rows[row].cells
        for cell in np.arange(5):
            rw[cell].text = wrds[row][cell]
            rw[cell].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            rw[cell].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

document = Document()

cardnum = int(input('Enter the number of Bingo cards to generate: '))
print('%s Bingo-cards will be created in 1 word-file' % cardnum)

for i in np.arange(cardnum):
    document.add_heading('PhD-Public-Defence-BINGO', 0)
    document.add_paragraph()
    p = document.add_paragraph(
        'Hello and welcome to the PhD-Public-Defence-BINGO. As the presentation takes place, simply cross out '
        'the words you hear that are on the bingo grid below. If you have a BINGO (one full line either horizontal, '
        'vertical or diagonal) please do ')
    p.add_run('NOT').bold = True
    p.add_run(' yell "Bingo". Instead, please note the time and the slide at which you had your Bingo.')

    document.add_paragraph()
    document.add_paragraph('Name:\t\t\t_________________________________________________________')
    document.add_paragraph('Bingo-time:\t\t_________________________________________________________')
    document.add_paragraph('Slide:\t\t\t_________________________________________________________')

    document.add_paragraph()
    create_bingogrid()

    if i < (cardnum - 1):
        document.add_page_break()


document.save('bingo.docx')
